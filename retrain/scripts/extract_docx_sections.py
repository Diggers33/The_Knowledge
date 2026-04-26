#!/usr/bin/env python3
"""Extract numbered sections from Part B DOCX files and append to the corpus.

Usage (from iris-kb/):
  python retrain/scripts/extract_docx_sections.py <file1.docx> [file2.docx ...]

Outputs:
  retrain/corpus/sections/<ACRONYM>__<section_id>.txt  (one file per section)
  Appends rows to retrain/corpus/inventory.csv
"""

import csv, re, sys
from pathlib import Path
import docx

SECTIONS_DIR = Path("retrain/corpus/sections")
INVENTORY    = Path("retrain/corpus/inventory.csv")

SECTION_HEADERS = [
    ("1.1", r"^\s*1\s*[.]\s*1\s+\w"),
    ("1.2", r"^\s*1\s*[.]\s*2\s+\w"),
    ("1.3", r"^\s*1\s*[.]\s*3\s+\w"),
    ("1.4", r"^\s*1\s*[.]\s*4\s+\w"),
    ("2.1", r"^\s*2\s*[.]\s*1\s+\w"),
    ("2.2", r"^\s*2\s*[.]\s*2\s+\w"),
    ("2.3", r"^\s*2\s*[.]\s*3\s+\w"),
    ("3.1", r"^\s*3\s*[.]\s*1\s+\w"),
    ("3.2", r"^\s*3\s*[.]\s*2\s+\w"),
    ("3.3", r"^\s*3\s*[.]\s*3\s+\w"),
    ("4",   r"^\s*4\s*[.]\s*(?!\d)\w"),
]

TOC_RE = re.compile(r"\.{5,}")

def get_acronym(stem: str) -> str:
    # e.g. PartB_VISUM-LRS_final → VISUM-LRS, PhotonX_PartB_upload2 → PHOTONX
    # Strip common PartB prefixes/suffixes
    s = re.sub(r"(?i)(partb|part_b|part-b|upload\d*|final|v\d+|_|-)", " ", stem)
    tokens = re.findall(r"[A-Z][A-Z0-9\-]{2,}", s.upper())
    if tokens:
        return tokens[0].strip("-")
    return stem[:12].upper()

def extract_docx_text(path: Path) -> str:
    doc = docx.Document(str(path))
    lines = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            lines.append("")
            continue
        # Promote heading paragraphs: ensure they appear on their own line
        style = (para.style.name or "").lower()
        if "heading" in style or re.match(r"^\d+\.\d*\s+[A-Z]", text):
            lines.append("")
            lines.append(text)
            lines.append("")
        else:
            lines.append(text)
    # Also pull text from tables
    for table in doc.tables:
        for row in table.rows:
            lines.append(" | ".join(cell.text.strip() for cell in row.cells if cell.text.strip()))
    return "\n".join(lines)

def split_sections(full: str) -> dict[str, str]:
    indices = []
    for sid, pat in SECTION_HEADERS:
        m = re.search(pat, full, flags=re.MULTILINE | re.IGNORECASE)
        if m:
            indices.append((sid, m.start()))
    indices.sort(key=lambda x: x[1])

    out: dict[str, str] = {}
    for i, (sid, start) in enumerate(indices):
        end = indices[i + 1][1] if i + 1 < len(indices) else len(full)
        text = full[start:end].strip()
        # Remove TOC lines
        text = "\n".join(l for l in text.splitlines() if not TOC_RE.search(l))
        text = re.sub(r"\n{3,}", "\n\n", text).strip()
        if len(text.split()) >= 200:
            out[sid] = text
    return out

def load_existing_inventory() -> list[dict]:
    if not INVENTORY.exists():
        return []
    return list(csv.DictReader(open(INVENTORY, encoding="utf-8")))

def append_inventory(rows: list[dict]):
    existing = load_existing_inventory()
    all_rows = existing + rows
    fieldnames = ["acronym", "grant", "call_id", "section_id", "word_count", "char_count", "source_pdf", "extracted_text"]
    with open(INVENTORY, "w", newline="", encoding="utf-8") as f:
        w = csv.DictWriter(f, fieldnames=fieldnames)
        w.writeheader()
        w.writerows(all_rows)

def main():
    if len(sys.argv) < 2:
        print("Usage: extract_docx_sections.py <file.docx> [...]")
        sys.exit(1)

    SECTIONS_DIR.mkdir(parents=True, exist_ok=True)

    # Load existing inventory to avoid duplicates
    existing = load_existing_inventory()
    existing_keys = {(r["acronym"], r["section_id"]) for r in existing}

    new_rows: list[dict] = []

    for docx_path_str in sys.argv[1:]:
        docx_path = Path(docx_path_str)
        if not docx_path.exists():
            print(f"  ERROR: not found: {docx_path}", file=sys.stderr)
            continue

        acronym = get_acronym(docx_path.stem)
        print(f"\n{acronym}  ({docx_path.name})")

        full = extract_docx_text(docx_path)
        sections = split_sections(full)

        if not sections:
            print("  -> no sections detected (check heading formatting)")
            continue

        grant_m = re.search(r"\b(1\d{8})\b", full)
        grant   = grant_m.group(1) if grant_m else ""
        call_m  = re.search(r"HORIZON[-_]CL\d[-_]20\d\d[-_][A-Z0-9\-]+", full, re.IGNORECASE)
        call    = call_m.group(0).upper() if call_m else ""

        for sid, text in sections.items():
            key = (acronym, sid)
            if key in existing_keys:
                print(f"  SKIP {sid}: already in inventory")
                continue

            safe_sid = sid.replace(".", "_")
            out_path = SECTIONS_DIR / f"{acronym}__{safe_sid}.txt"
            out_path.write_text(text, encoding="utf-8")

            wc = len(text.split())
            new_rows.append({
                "acronym":        acronym,
                "grant":          grant,
                "call_id":        call,
                "section_id":     sid,
                "word_count":     wc,
                "char_count":     len(text),
                "source_pdf":     docx_path.name,
                "extracted_text": str(out_path),
            })
            print(f"  -> {sid}: {wc} words  [{out_path.name}]")

        print(f"  {len([s for s in sections if (acronym, s) not in existing_keys])} new sections extracted")

    if new_rows:
        append_inventory(new_rows)
        print(f"\nAppended {len(new_rows)} rows to {INVENTORY}")
    else:
        print("\nNo new rows added.")

if __name__ == "__main__":
    main()
